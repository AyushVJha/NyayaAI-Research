"""
Document Ingestion Pipeline for Indian Legal Documents
Handles multiple file formats: PDF, DOCX, DOC, TXT, and images
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
import mimetypes

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
# import textract  # Commented out due to dependency conflict
from PIL import Image
import fitz  # PyMuPDF

from config.settings import DOCUMENT_CONFIG
from utils.logger import logger
from utils.error_handler import handle_exceptions, DocumentProcessingError, ValidationError
from utils.text_normalization import normalize_legal_text, extract_document_sections
from ingestion.ocr_pipeline import IndianLegalOCR

class LegalDocumentProcessor:
    """
    Comprehensive document processor for Indian legal documents
    Supports multiple formats and extraction methods
    """
    
    def __init__(self):
        self.supported_formats = DOCUMENT_CONFIG["supported_formats"]
        self.max_file_size_mb = DOCUMENT_CONFIG["max_file_size_mb"]
        self.chunk_size = DOCUMENT_CONFIG["chunk_size"]
        self.overlap = DOCUMENT_CONFIG["overlap"]
        self.ocr_processor = IndianLegalOCR()
        
    @handle_exceptions(DocumentProcessingError, "Document processing failed")
    def process_document(self, file_path: str, 
                        normalize_text: bool = True,
                        extract_sections: bool = True,
                        ocr_fallback: bool = True) -> Dict[str, Any]:
        """
        Process a legal document and extract text content
        
        Args:
            file_path: Path to the document file
            normalize_text: Whether to apply text normalization
            extract_sections: Whether to extract document sections
            ocr_fallback: Whether to use OCR if text extraction fails
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info(f"Processing document: {file_path}")
        
        # Validate file
        self._validate_file(file_path)
        
        # Determine file type and processing method
        file_extension = Path(file_path).suffix.lower()
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        
        # Extract text based on file type
        extraction_result = self._extract_text_by_type(file_path, file_extension, ocr_fallback)
        
        # Post-process extracted text
        processed_result = self._post_process_text(
            extraction_result, 
            normalize_text, 
            extract_sections
        )
        
        # Add metadata
        processed_result.update({
            'file_info': {
                'filename': Path(file_path).name,
                'file_path': file_path,
                'file_size_mb': round(file_size_mb, 2),
                'file_type': file_extension,
                'mime_type': mimetypes.guess_type(file_path)[0]
            },
            'processing_info': {
                'normalize_text': normalize_text,
                'extract_sections': extract_sections,
                'ocr_fallback': ocr_fallback,
                'extraction_method': extraction_result.get('method', 'unknown')
            }
        })
        
        logger.info(f"Document processing completed. Extracted {len(processed_result['text'])} characters")
        return processed_result
    
    def _validate_file(self, file_path: str):
        """Validate file before processing"""
        if not os.path.exists(file_path):
            raise ValidationError(f"File not found: {file_path}", "FILE_NOT_FOUND")
        
        file_extension = Path(file_path).suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValidationError(
                f"Unsupported file format: {file_extension}. Supported formats: {self.supported_formats}",
                "UNSUPPORTED_FORMAT"
            )
        
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValidationError(
                f"File too large: {file_size_mb:.2f}MB. Maximum size: {self.max_file_size_mb}MB",
                "FILE_TOO_LARGE"
            )
    
    def _extract_text_by_type(self, file_path: str, file_extension: str, ocr_fallback: bool) -> Dict[str, Any]:
        """Extract text based on file type"""
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path, ocr_fallback)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_from_word(file_path)
        elif file_extension == '.txt':
            return self._extract_from_text(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_from_image(file_path)
        else:
            # Try textract as fallback
            return self._extract_with_textract(file_path)
    
    def _extract_from_pdf(self, file_path: str, ocr_fallback: bool) -> Dict[str, Any]:
        """Extract text from PDF files"""
        logger.debug(f"Extracting text from PDF: {file_path}")
        
        text_content = ""
        method = "pdf_text_extraction"
        metadata = {}
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(f"--- Page {i+1} ---\n{page_text}")
                
                text_content = "\n\n".join(pages_text)
                metadata = {
                    'total_pages': len(pdf.pages),
                    'pdf_info': pdf.metadata if hasattr(pdf, 'metadata') else {}
                }
            
            # If no text extracted and OCR fallback enabled
            if (not text_content.strip() or len(text_content.strip()) < 50) and ocr_fallback:
                logger.info("PDF text extraction yielded minimal content, trying OCR")
                return self._extract_pdf_with_ocr(file_path)
            
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}, trying PyPDF2")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages_text = []
                    
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(f"--- Page {i+1} ---\n{page_text}")
                    
                    text_content = "\n\n".join(pages_text)
                    metadata = {
                        'total_pages': len(pdf_reader.pages),
                        'pdf_info': pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
                    }
                    method = "pypdf2_extraction"
                    
            except Exception as e2:
                if ocr_fallback:
                    logger.info("PDF text extraction failed, using OCR")
                    return self._extract_pdf_with_ocr(file_path)
                else:
                    raise DocumentProcessingError(
                        f"PDF text extraction failed: {str(e2)}",
                        "PDF_EXTRACTION_FAILED"
                    )
        
        return {
            'text': text_content,
            'method': method,
            'metadata': metadata,
            'success': True
        }
    
    def _extract_pdf_with_ocr(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using OCR"""
        logger.info("Using OCR for PDF text extraction")
        
        try:
            # Convert PDF pages to images
            pdf_document = fitz.open(file_path)
            pages_text = []
            
            for page_num in range(len(pdf_document)):
                # Convert page to image
                page = pdf_document[page_num]
                mat = fitz.Matrix(2.0, 2.0)  # Increase resolution
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Process with OCR
                ocr_result = self.ocr_processor.process_image(img_data)
                if ocr_result['processing_successful']:
                    pages_text.append(f"--- Page {page_num+1} ---\n{ocr_result['text']}")
            
            pdf_document.close()
            
            return {
                'text': "\n\n".join(pages_text),
                'method': "pdf_ocr",
                'metadata': {'total_pages': len(pdf_document)},
                'success': True
            }
            
        except Exception as e:
            raise DocumentProcessingError(
                f"PDF OCR extraction failed: {str(e)}",
                "PDF_OCR_FAILED"
            )
    
    def _extract_from_word(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word documents"""
        logger.debug(f"Extracting text from Word document: {file_path}")
        
        try:
            if file_path.endswith('.docx'):
                # Use python-docx for .docx files
                doc = Document(file_path)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                text_content = "\n\n".join(paragraphs)
                
                # Extract metadata
                metadata = {
                    'paragraphs_count': len(doc.paragraphs),
                    'core_properties': {
                        'title': doc.core_properties.title,
                        'author': doc.core_properties.author,
                        'subject': doc.core_properties.subject,
                        'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                        'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None
                    }
                }
                method = "python_docx"
                
            else:
                # Fallback for .doc files - use PyMuPDF or raise error
                raise DocumentProcessingError(
                    "Legacy .doc files are not supported. Please convert to .docx format.",
                    "DOC_FORMAT_NOT_SUPPORTED"
                )
            
            return {
                'text': text_content,
                'method': method,
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Word document extraction failed: {str(e)}",
                "WORD_EXTRACTION_FAILED"
            )
    
    def _extract_from_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files"""
        logger.debug(f"Reading text file: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    
                    return {
                        'text': text_content,
                        'method': f"text_file_{encoding}",
                        'metadata': {'encoding': encoding, 'file_size': len(text_content)},
                        'success': True
                    }
                    
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError(
                "Could not decode text file with any supported encoding",
                "TEXT_ENCODING_FAILED"
            )
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Text file reading failed: {str(e)}",
                "TEXT_FILE_FAILED"
            )
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image files using OCR"""
        logger.debug(f"Extracting text from image: {file_path}")
        
        try:
            ocr_result = self.ocr_processor.process_image(file_path)
            
            return {
                'text': ocr_result['text'],
                'method': "image_ocr",
                'metadata': {
                    'ocr_confidence': ocr_result['statistics']['avg_confidence'],
                    'image_size': ocr_result['statistics']['original_image_size'],
                    'languages_used': ocr_result['languages_used'],
                    'quality_issues': ocr_result['quality_issues']
                },
                'success': ocr_result['processing_successful']
            }
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Image OCR failed: {str(e)}",
                "IMAGE_OCR_FAILED"
            )
    
    def _extract_with_textract(self, file_path: str) -> Dict[str, Any]:
        """Extract text using alternative methods (textract removed due to dependency conflict)"""
        logger.debug(f"Using alternative extraction for: {file_path}")
        
        # Try PyMuPDF as fallback for various formats
        try:
            doc = fitz.open(file_path)
            text_content = ""
            for page in doc:
                text_content += page.get_text()
            doc.close()
            
            return {
                'text': text_content,
                'method': "pymupdf_fallback",
                'metadata': {},
                'success': True
            }
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Alternative extraction failed: {str(e)}",
                "ALTERNATIVE_EXTRACTION_FAILED"
            )
    
    def _post_process_text(self, extraction_result: Dict[str, Any], 
                          normalize_text: bool, 
                          extract_sections: bool) -> Dict[str, Any]:
        """Post-process extracted text"""
        
        text = extraction_result['text']
        
        # Apply text normalization if requested
        if normalize_text and text:
            try:
                normalized_text = normalize_legal_text(text, preserve_structure=True)
                extraction_result['normalized_text'] = normalized_text
            except Exception as e:
                logger.warning(f"Text normalization failed: {str(e)}")
                extraction_result['normalized_text'] = text
        
        # Extract document sections if requested
        if extract_sections and text:
            try:
                sections = extract_document_sections(text)
                extraction_result['sections'] = sections
            except Exception as e:
                logger.warning(f"Section extraction failed: {str(e)}")
                extraction_result['sections'] = {}
        
        # Add text statistics
        extraction_result['text_statistics'] = {
            'character_count': len(text),
            'word_count': len(text.split()) if text else 0,
            'line_count': len(text.splitlines()) if text else 0,
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]) if text else 0
        }
        
        return extraction_result
    
    def process_multiple_documents(self, file_paths: List[str], 
                                 combine_results: bool = False) -> Dict[str, Any]:
        """Process multiple documents"""
        logger.info(f"Processing {len(file_paths)} documents")
        
        results = []
        combined_text = []
        total_characters = 0
        successful_count = 0
        
        for i, file_path in enumerate(file_paths):
            try:
                logger.info(f"Processing document {i+1}/{len(file_paths)}: {Path(file_path).name}")
                result = self.process_document(file_path)
                
                results.append(result)
                if result.get('success', False):
                    successful_count += 1
                    total_characters += len(result['text'])
                    
                    if combine_results:
                        combined_text.append(f"=== Document: {Path(file_path).name} ===\n{result['text']}")
                
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {str(e)}")
                results.append({
                    'file_info': {'filename': Path(file_path).name, 'file_path': file_path},
                    'error': str(e),
                    'success': False
                })
        
        batch_result = {
            'individual_results': results,
            'batch_statistics': {
                'total_documents': len(file_paths),
                'successful_documents': successful_count,
                'failed_documents': len(file_paths) - successful_count,
                'total_characters': total_characters,
                'success_rate': successful_count / len(file_paths) if file_paths else 0
            }
        }
        
        if combine_results:
            batch_result['combined_text'] = '\n\n'.join(combined_text)
        
        return batch_result

# Convenience functions
def extract_text_from_document(file_path: str, normalize: bool = True) -> str:
    """Simple function to extract text from a document"""
    processor = LegalDocumentProcessor()
    result = processor.process_document(file_path, normalize_text=normalize)
    return result.get('normalized_text', result['text'])

def process_legal_document_batch(file_paths: List[str]) -> Dict[str, Any]:
    """Process multiple legal documents"""
    processor = LegalDocumentProcessor()
    return processor.process_multiple_documents(file_paths, combine_results=True)

def get_document_sections(file_path: str) -> Dict[str, str]:
    """Extract sections from a legal document"""
    processor = LegalDocumentProcessor()
    result = processor.process_document(file_path, extract_sections=True)
    return result.get('sections', {})
